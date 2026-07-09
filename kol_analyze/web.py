"""本地 Web 应用（给 2-3 人用）。

    python -m kol_analyze.web         # 浏览器打开 http://127.0.0.1:8000

流程：选产品(RM/RC/RO) → 上传后台 xlsx + 大盘截图 → 审阅并修正命名(存记忆库)
     → 补充缺失 → 生成复盘 docx 并下载。每次生成都归档进「历史复盘」。

生成分析用 Claude Code 订阅（CLI），无需官方 API key。
记忆库与历史按产品分开存。
"""

from __future__ import annotations

import os
import threading
import uuid
from datetime import datetime, timezone
from pathlib import Path

from flask import (Flask, jsonify, redirect, request, send_file, session)

from . import (analyzer, docx_writer, engine, loader, market, memory,
               metrics, scripts, staffing, store, vision)
from .config import PRODUCTS, Settings
from .web_ui import LOGIN_PAGE, PAGE

app = Flask(__name__)
# 部署到线上时用固定密钥（环境变量），本地随机即可
app.secret_key = os.environ.get("KOL_SECRET") or ("kol-local-" + uuid.uuid4().hex)
app.json.sort_keys = False  # 保持产品插入顺序（RM/RC/RO），别按字母排

# 访问密码：设了 KOL_PASSWORD 就要求登录（线上必设，防止链接被外人打开）
ACCESS_PW = os.environ.get("KOL_PASSWORD", "")


@app.before_request
def _guard():
    if not ACCESS_PW:
        return None
    if request.path in ("/login", "/favicon.ico"):
        return None
    if session.get("authed"):
        return None
    if request.path.startswith("/api/"):
        return jsonify({"ok": False, "error": "未登录"}), 401
    return redirect("/login")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        if request.form.get("pw") == ACCESS_PW:
            session["authed"] = True
            return redirect("/")
        return LOGIN_PAGE.replace("<!--ERR-->", "密码不对，再试一次。")
    return LOGIN_PAGE.replace("<!--ERR-->", "")

SESSIONS: dict[str, dict] = {}
SETTINGS = Settings()

# 产品清单：默认 RM/RC/RO，可在页面自定义添加，存在持久盘，重启不丢
_PRODUCT_SEED = dict(PRODUCTS)
PRODUCTS.clear()
PRODUCTS.update(store.load_products(_PRODUCT_SEED))
DEFAULT_PRODUCT = next(iter(PRODUCTS))


def _sid() -> str:
    if "sid" not in session:
        session["sid"] = uuid.uuid4().hex
    sid = session["sid"]
    if sid not in SESSIONS:
        SESSIONS[sid] = {"sid": sid, "product": DEFAULT_PRODUCT,
                         "market": market.MarketContext(),
                         "gen": {"status": "idle"}, "meta": {}}
    return sid


def _S() -> dict:
    return SESSIONS[_sid()]


def _wd(st: dict) -> Path:
    return store.session_dir(st["product"], st["sid"])


def _now() -> str:
    return datetime.now(timezone.utc).astimezone().strftime("%Y%m%d-%H%M")


def _save_project(st: dict, status=None, with_files=False, with_docx=False) -> str | None:
    """把当前会话存/更新成一个项目（草稿或已完成），返回项目 id。"""
    if "analysis" not in st:
        return None
    docx = None
    if with_docx and st.get("gen", {}).get("file"):
        docx = Path(st["gen"]["file"])
    stats = {"creatives": sum(l.count for l in st["analysis"].langs),
             "langs": len(st["analysis"].langs),
             "migrations": len(st["sa"].migrations) if st.get("sa") else 0}
    pid = store.save_draft(
        st["product"], _now(), st.get("meta", {}), market.to_dict(st["market"]),
        st.get("data"), _wd(st) / "data", did=st.get("project_id"),
        status=status, stats=stats, staffing=st.get("staffing", ""),
        extra=st.get("extra", ""), copy_files=with_files, docx_path=docx)
    st["project_id"] = pid
    return pid


# --------------------------------------------------------------------------
# 核心：跑分析（不含 Claude 写作，秒级）
# --------------------------------------------------------------------------

def _recompute(st: dict) -> dict:
    mem = memory.load(store.memory_path(st["product"]))
    data_dir = _wd(st) / "data"
    files = [f for f in data_dir.iterdir()
             if f.suffix.lower() in (".xlsx", ".xls", ".csv")]
    if not files:
        return {"ok": False, "error": "还没有上传数据文件。"}

    ds = loader.load(data_dir, mem)
    mkt = st["market"]
    if ds.kol_total_spend and not mkt.kol_total_spend:
        mkt.kol_total_spend = ds.kol_total_spend

    analysis = metrics.compute(ds, mkt, SETTINGS.thresholds,
                               assume_complete=bool(st.get("force_complete")))
    langs_sorted = sorted(analysis.langs, key=lambda l: l.spend_share, reverse=True)
    top = langs_sorted[0].spend_share if langs_sorted else 0.0
    dominant = (langs_sorted[0].name if langs_sorted else "")
    # 你确认过数据完整时（force_complete），就按真实情况判定、不再标「待补全」
    if st.get("force_complete"):
        incomplete = set()
    else:
        incomplete = ({l.lang for l in analysis.langs if l.spend_share < 5.0}
                      if top >= 85.0 else set())
    sa = scripts.analyze(analysis.langs, SETTINGS.thresholds, incomplete, mem)

    st["staffing"] = store.load_staffing(st["product"])
    st["sop"] = store.load_sop(st["product"])
    st.update(ds=ds, mem=mem, analysis=analysis, sa=sa, incomplete=incomplete,
              dominant=dominant, dominant_share=round(top, 1))
    return {"ok": True, **_snapshot(st)}


def _snapshot(st: dict) -> dict:
    a, sa, mem = st["analysis"], st["sa"], st["mem"]
    rows = []
    for l in a.langs:
        for c in l.creatives:
            rows.append({
                "ad_name": c.ad_name, "play": c.play or "", "lang": l.name,
                "influencer": c.influencer or "",
                "scripts": scripts.script_themes(c, mem),
                "formats": scripts.format_tags(c, mem), "tier": c.tier,
                "spend": round(c.spend, 1), "roi7": round((c.roi7 or 0) * 100, 1),
            })
    rows.sort(key=lambda r: r["spend"], reverse=True)
    gaps = [{"name": g.name, "ad": g.ad_market_share, "kol": g.kol_spend_share,
             "pub": g.publish_share, "breakout": g.breakout_rate,
             "verdict": g.verdict, "line": g.one_line} for g in a.gaps]
    migrations = [{"theme": r.theme, "best": r.best_lang, "to": r.migrate_to,
                   "reason": r.reason} for r in sa.migrations[:12]]
    return {
        "product": st["product"], "product_name": PRODUCTS.get(st["product"], st["product"]),
        "detected_products": st["ds"].products,
        "stats": {"creatives": sum(l.count for l in a.langs),
                  "langs": len(a.langs), "scripts": len(sa.scripts),
                  "migrations": len(sa.migrations),
                  "kol_share": st["market"].kol_share_of_total},
        "langs": [{"name": l.name, "count": l.count,
                   "spend_share": round(l.spend_share, 1),
                   "breakout": round(l.breakout_rate, 1) if l.breakout_rate is not None else None}
                  for l in a.langs],
        "rows": rows, "gaps": gaps, "migrations": migrations,
        "memory": memory.to_dict(mem), "missing": _missing(st),
        "notes": st["market"].notes, "history": _history_list(st["product"]),
        "staffing": st.get("staffing", ""), "sop": st.get("sop", ""),
        "extra": st.get("extra", ""),
    }


def _start_vision(st: dict, shots: list):
    """后台识别大盘截图，读到就更新 market 并重算，不阻塞页面。"""
    st["vision"] = {"status": "running"}

    def worker(state):
        try:
            mkt = vision.read_screenshots(shots, SETTINGS)
            if not mkt.is_empty():
                state["market"] = mkt
                snap = _recompute(state)
                state["vision"] = {"status": "done", "snapshot": snap}
            else:
                state["vision"] = {"status": "empty",
                                   "note": "；".join(mkt.notes) or "没读出占比"}
        except Exception as e:  # noqa
            state["vision"] = {"status": "failed", "note": str(e)}

    threading.Thread(target=worker, args=(st,), daemon=True).start()


@app.get("/api/vision/status")
def api_vision_status():
    return jsonify(_S().get("vision", {"status": "idle"}))


def _missing(st: dict) -> dict:
    mkt = st["market"]
    inc = [l.name for l in st["analysis"].langs if l.lang in st.get("incomplete", set())]
    return {"market": mkt.is_empty(), "incomplete_langs": sorted(inc),
            "dominant": st.get("dominant", ""),
            "dominant_share": st.get("dominant_share"),
            "force_complete": bool(st.get("force_complete")),
            "coverage_gaps": st["analysis"].coverage_gaps}


def _history_list(product: str) -> list[dict]:
    return [{"id": e.id, "period": e.period, "title": e.title,
             "created": e.created, "stats": e.stats}
            for e in store.list_history(product)]


# --------------------------------------------------------------------------
# 路由
# --------------------------------------------------------------------------

@app.get("/")
def index():
    _sid()
    return PAGE


@app.get("/api/engine")
def api_engine():
    return jsonify({"engine": engine.available(), "label": engine.label(),
                    "products": PRODUCTS})


@app.post("/api/analyze")
def api_analyze():
    st = _S()
    st["product"] = request.form.get("product") or DEFAULT_PRODUCT
    st["meta"] = {"title": request.form.get("title") or "月度 KOL 广告复盘",
                  "period": request.form.get("period") or ""}
    wd = _wd(st)
    for f in request.files.getlist("data"):
        if f.filename:
            f.save(wd / "data" / Path(f.filename).name)
    shots = []
    for f in request.files.getlist("shots"):
        if f.filename:
            p = wd / "shots" / Path(f.filename).name
            f.save(p)
            shots.append(str(p))
    st["project_id"] = None  # 新上传 = 新项目
    snap = _recompute(st)
    if snap.get("ok"):
        _save_project(st, status="draft", with_files=True)
    if shots:
        _start_vision(st, shots)
        snap["vision"] = "running"
    return jsonify(snap)


@app.post("/api/supplement")
def api_supplement():
    st = _S()
    wd = _wd(st)
    for f in request.files.getlist("data"):
        if f.filename:
            f.save(wd / "data" / Path(f.filename).name)
    shots = []
    for f in request.files.getlist("shots"):
        if f.filename:
            p = wd / "shots" / Path(f.filename).name
            f.save(p)
            shots.append(str(p))
    snap = _recompute(st)
    if shots:
        _start_vision(st, shots)
        snap["vision"] = "running"
    return jsonify(snap)


@app.post("/api/market")
def api_market():
    st = _S()
    st["market"] = market.from_dict(request.get_json(force=True))
    return jsonify(_recompute(st))


@app.get("/api/staffing")
def api_staffing_get():
    prod = request.args.get("product") or _S()["product"]
    return jsonify({"product": prod, "text": store.load_staffing(prod)})


@app.post("/api/staffing")
def api_staffing_set():
    st = _S()
    text = (request.get_json(silent=True) or {}).get("text", "")
    store.save_staffing(st["product"], text)
    st["staffing"] = text
    return jsonify(_recompute(st) if "analysis" in st else {"ok": True})


@app.get("/api/sop")
def api_sop_get():
    prod = request.args.get("product") or _S()["product"]
    return jsonify({"product": prod, "text": store.load_sop(prod)})


@app.post("/api/sop")
def api_sop_set():
    st = _S()
    text = (request.get_json(silent=True) or {}).get("text", "")
    store.save_sop(st["product"], text)
    st["sop"] = text
    return jsonify(_recompute(st) if "analysis" in st else {"ok": True})


@app.post("/api/sop/upload")
def api_sop_upload():
    """上传 SOP 文档（docx/txt/md），抽取文字存为本产品 SOP。"""
    st = _S()
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "没有文件"})
    text = _extract_doc_text(f)
    if not text:
        return jsonify({"ok": False, "error": "没读到文字（支持 .docx/.txt/.md）"})
    store.save_sop(st["product"], text)
    st["sop"] = text
    return jsonify({"ok": True, "text": text})


def _extract_doc_text(f) -> str:
    name = (f.filename or "").lower()
    if name.endswith((".txt", ".md", ".csv")):
        return f.read().decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        try:
            import docx
            d = docx.Document(f)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""
    return ""


@app.post("/api/extra")
def api_extra():
    st = _S()
    st["extra"] = (request.get_json(silent=True) or {}).get("text", "")
    return jsonify({"ok": True})


@app.post("/api/complete")
def api_complete():
    """用户确认「这些语言本月就是没怎么投放，数据是完整的」-> 按真实情况判定。"""
    st = _S()
    st["force_complete"] = bool((request.get_json(silent=True) or {}).get("value", True))
    return jsonify(_recompute(st))


@app.post("/api/correct")
def api_correct():
    st = _S()
    body = request.get_json(force=True)
    mem = memory.load(store.memory_path(st["product"]))
    scope = body.get("scope", "keyword")
    common = dict(add_script=body.get("add_script", []),
                  set_script=body.get("set_script", []),
                  add_format=body.get("add_format", []), note=body.get("note", ""))
    if scope == "exact":
        mem.add_play_override(body.get("key", ""), **common)
    elif scope == "keyword":
        mem.keyword_rules.append(memory.KeywordRule(match=body.get("match", ""), **common))
    elif scope == "influencer":
        mem.influencer_alias[body.get("from", "")] = body.get("to", "")
    elif scope == "lang":
        mem.lang_overrides[body.get("key", "")] = body.get("to", "")
    mem.save(store.memory_path(st["product"]))
    return jsonify(_recompute(st))


@app.post("/api/generate")
def api_generate():
    st = _S()
    if "analysis" not in st:
        return jsonify({"ok": False, "error": "请先上传并分析数据。"})
    if st["gen"]["status"] == "running":
        return jsonify({"ok": True, "status": "running"})
    st["gen"] = {"status": "running", "engine": engine.label()}

    def worker(state):
        try:
            sf = staffing.build_facts(state.get("staffing", ""),
                                      state["analysis"].gaps, state["sa"].lang_strategies)
            data = analyzer.analyze(state["analysis"], state["sa"], SETTINGS,
                                    state["meta"]["title"], state["meta"]["period"],
                                    mem=state.get("mem"), staffing_facts=sf,
                                    sop=state.get("sop", ""), extra=state.get("extra", ""))
            out = _wd(state) / "复盘.docx"
            docx_writer.render(data, state["analysis"], state["sa"], out)
            created = datetime.now(timezone.utc).astimezone().strftime("%Y%m%d-%H%M")
            stats = {"creatives": sum(l.count for l in state["analysis"].langs),
                     "langs": len(state["analysis"].langs),
                     "migrations": len(state["sa"].migrations)}
            entry = store.archive(state["product"], state["meta"]["period"] or "复盘",
                                  state["meta"]["title"], created, out, data, stats)
            state["data"] = data
            state["gen"] = {"status": "done", "file": str(out), "hid": entry.id}
            _save_project(state, status="completed", with_docx=True)
        except Exception as e:  # noqa
            state["gen"] = {"status": "error", "error": str(e)}

    threading.Thread(target=worker, args=(st,), daemon=True).start()
    return jsonify({"ok": True, "status": "running", "engine": engine.label()})


@app.get("/api/status")
def api_status():
    return jsonify(_S()["gen"])


# ---- 复盘预览 · 分块修订 ----

_LIST_BLOCKS = {"script_section.migrations", "script_section.format_suggestions"}


def _get_path(data, path):
    cur = data
    for p in path.split("."):
        if isinstance(cur, list):
            if not p.isdigit() or int(p) >= len(cur):
                return None
            cur = cur[int(p)]
        elif isinstance(cur, dict):
            cur = cur.get(p)
        else:
            return None
        if cur is None:
            return None
    return cur


def _set_path(data, path, val):
    parts = path.split(".")
    cur = data
    for p in parts[:-1]:
        cur = cur[int(p)] if isinstance(cur, list) else cur.setdefault(p, {})
    last = parts[-1]
    if isinstance(cur, list):
        cur[int(last)] = val
    else:
        cur[last] = val


def _apply_block(data, key, text):
    if key in _LIST_BLOCKS:
        _set_path(data, key, [x for x in text.split("\n") if x.strip()])
    else:
        _set_path(data, key, text)


def _blocks(data):
    out = []

    def add(path, label):
        v = _get_path(data, path)
        if v is None:
            return
        text = "\n".join(v) if (path in _LIST_BLOCKS and isinstance(v, list)) else str(v)
        out.append({"key": path, "label": label, "text": text,
                    "list": path in _LIST_BLOCKS})

    add("ad_section.overview", "一、广告部份 · 概述")
    add("ad_section.caveat", "口径提醒")
    add("gap_summary", "二、KOL 分语言 · 一句话总结")
    add("script_section.overview", "三、素材/脚本 · 概述")
    add("script_section.migrations", "跨语言迁移建议（每行一条）")
    add("script_section.format_suggestions", "形式覆盖建议（每行一条）")
    for i, l in enumerate(data.get("langs", [])):
        nm = l.get("name", f"语言{i}")
        add(f"langs.{i}.one_liner", f"{nm} · 定位")
        add(f"langs.{i}.conversion", f"{nm} · 转化情况")
        add(f"langs.{i}.creative_analysis", f"{nm} · 素材分析")
        add(f"langs.{i}.todo", f"{nm} · todo")
    for i, s in enumerate(data.get("script_section", {}).get("lang_strategies", [])):
        nm = s.get("name", f"语言{i}")
        add(f"script_section.lang_strategies.{i}.suggestion", f"{nm} · 脚本策略")
    add("staffing_section.overview", "四、人力分工 · 总述")
    for i, p in enumerate(data.get("staffing_section", {}).get("people", [])):
        nm = p.get("person", f"成员{i}")
        add(f"staffing_section.people.{i}.suggestion", f"{nm} · 调整建议")
    return out


def _rerender(st: dict):
    """把当前（可能已修订的）data 重渲染到工作 docx，并覆盖历史归档。"""
    out = _wd(st) / "复盘.docx"
    docx_writer.render(st["data"], st["analysis"], st["sa"], out)
    st["gen"]["file"] = str(out)
    hid = st["gen"].get("hid")
    if hid:
        store.update_entry(st["product"], hid, out, st["data"])


@app.get("/api/report")
def api_report():
    st = _S()
    if "data" not in st:
        return jsonify({"ok": False, "error": "还没有生成结果。"})
    return jsonify({"ok": True, "blocks": _blocks(st["data"]),
                    "can_ai": engine.available() != "offline",
                    "label": st["meta"].get("title", "")})


@app.post("/api/revise")
def api_revise():
    st = _S()
    if "data" not in st:
        return jsonify({"ok": False, "error": "还没有生成结果。"})
    body = request.get_json(force=True)
    key = body.get("key", "")
    instruction = body.get("instruction", "").strip()
    reason = body.get("reason", "").strip()
    if not instruction:
        return jsonify({"ok": False, "error": "请填写想改成什么样。"})
    old = _get_path(st["data"], key)
    old_text = "\n".join(old) if isinstance(old, list) else str(old or "")
    label = next((b["label"] for b in _blocks(st["data"]) if b["key"] == key), key)

    new_text = analyzer.revise_passage(label, old_text, instruction, reason, SETTINGS)
    if not new_text:
        return jsonify({"ok": False, "error": "AI 修订不可用或失败（可手动改）。"})

    _apply_block(st["data"], key, new_text)
    # 记入记忆库（写作偏好，按产品）
    mem = memory.load(store.memory_path(st["product"]))
    created = datetime.now(timezone.utc).astimezone().strftime("%Y%m%d-%H%M")
    mem.add_style_note(scope=label, instruction=instruction, reason=reason, created=created)
    mem.save(store.memory_path(st["product"]))
    st["mem"] = mem
    _rerender(st)
    return jsonify({"ok": True, "key": key, "text": new_text,
                    "list": key in _LIST_BLOCKS,
                    "style_notes": len(mem.style_notes)})


@app.post("/api/polish")
def api_polish():
    """一键润色全文：先并入手动编辑，再让 Claude 统一润色，重渲染。"""
    st = _S()
    if "data" not in st:
        return jsonify({"ok": False, "error": "还没有生成结果。"})
    body = request.get_json(silent=True) or {}
    for b in body.get("blocks", []):
        _apply_block(st["data"], b["key"], b.get("text", ""))
    blocks_map = {b["key"]: b["text"] for b in _blocks(st["data"])}
    polished = analyzer.polish_document(blocks_map, SETTINGS, st.get("mem"))
    if not polished:
        return jsonify({"ok": False, "error": "润色不可用或失败（可手动改）。"})
    for k, v in polished.items():
        if isinstance(v, str) and _get_path(st["data"], k) is not None:
            _apply_block(st["data"], k, v)
    _rerender(st)
    return jsonify({"ok": True, "blocks": _blocks(st["data"])})


@app.post("/api/report_save")
def api_report_save():
    """保存手动修改（可选记入记忆库），重渲染 docx。"""
    st = _S()
    if "data" not in st:
        return jsonify({"ok": False, "error": "还没有生成结果。"})
    body = request.get_json(force=True)
    remember = []
    for b in body.get("blocks", []):
        _apply_block(st["data"], b["key"], b.get("text", ""))
        if b.get("remember") and b.get("note"):
            remember.append((b["key"], b["note"]))
    if remember:
        mem = memory.load(store.memory_path(st["product"]))
        created = datetime.now(timezone.utc).astimezone().strftime("%Y%m%d-%H%M")
        for key, note in remember:
            label = next((x["label"] for x in _blocks(st["data"]) if x["key"] == key), key)
            mem.add_style_note(scope=label, instruction=note, created=created)
        mem.save(store.memory_path(st["product"]))
        st["mem"] = mem
    _rerender(st)
    return jsonify({"ok": True})


@app.get("/api/download")
def api_download():
    st = _S()
    g = st["gen"]
    if g.get("status") != "done":
        return jsonify({"ok": False, "error": "还没生成完成。"}), 400
    title = st["meta"].get("title", "复盘")
    period = st["meta"].get("period", "")
    return send_file(g["file"], as_attachment=True,
                     download_name=f"{title}_{period or '复盘'}.docx")


@app.post("/api/draft/save")
def api_draft_save():
    """保存草稿：上传的文件 + 大盘 + 标题周期 +（若已生成）当前编辑内容。"""
    st = _S()
    if "analysis" not in st:
        return jsonify({"ok": False, "error": "先上传并分析后再存草稿。"})
    did = _save_project(st, with_files=True,
                        with_docx=st.get("gen", {}).get("status") == "done")
    return jsonify({"ok": True, "id": did})


@app.get("/api/drafts")
def api_drafts():
    prod = request.args.get("product") or _S()["product"]
    return jsonify({"product": prod, "drafts": store.list_drafts(prod)})


@app.post("/api/product/add")
def api_product_add():
    body = request.get_json(force=True)
    updated = store.add_product(body.get("code", ""), body.get("name", ""), _PRODUCT_SEED)
    PRODUCTS.clear()
    PRODUCTS.update(updated)
    return jsonify({"ok": True, "products": PRODUCTS})


@app.get("/api/projects")
def api_projects():
    """首页画廊：跨所有产品列出项目（草稿/已完成）。"""
    _sid()
    return jsonify({"projects": store.list_all_projects(list(PRODUCTS)),
                    "products": PRODUCTS})


@app.get("/api/project/download")
def api_project_download():
    prod = request.args.get("product", "")
    dr = store.get_draft(prod, request.args.get("id", ""))
    if not dr or not dr.get("_report") or not Path(dr["_report"]).exists():
        return jsonify({"ok": False, "error": "该项目还没有生成的文档。"}), 404
    meta = dr.get("meta", {})
    return send_file(dr["_report"], as_attachment=True,
                     download_name=f"{meta.get('title', '复盘')}_{meta.get('period', '')}.docx")


@app.post("/api/draft/delete")
def api_draft_delete():
    st = _S()
    body = request.get_json(silent=True) or {}
    prod = body.get("product") or st["product"]
    store.delete_draft(prod, body.get("id", ""))
    return jsonify({"ok": True})


@app.post("/api/draft/resume")
def api_draft_resume():
    st = _S()
    body = request.get_json(force=True)
    prod = body.get("product") or st["product"]
    dr = store.get_draft(prod, body.get("id", ""))
    if not dr:
        return jsonify({"ok": False, "error": "草稿不存在。"})
    st["product"] = prod
    st["project_id"] = dr.get("id")  # 续改同一项目，不新建
    st["meta"] = dr.get("meta", {})
    st["market"] = market.from_dict(dr.get("market", {}))
    st["extra"] = dr.get("extra", "")
    if dr.get("staffing"):
        store.save_staffing(prod, dr["staffing"])
    st["force_complete"] = False
    # 把项目文件复制进本会话数据目录（先清空）
    ddir = _wd(st) / "data"
    for f in ddir.iterdir():
        if f.is_file():
            f.unlink()
    src = dr.get("_data_dir")
    if src and src.exists():
        import shutil as _sh
        for f in src.iterdir():
            if f.is_file():
                _sh.copy(f, ddir / f.name)
    snap = _recompute(st)
    has_data = bool(dr.get("data"))
    if has_data:
        st["data"] = dr["data"]
        st["gen"] = {"status": "done"}
        _rerender(st)
    return jsonify({"ok": True, "has_data": has_data, **snap})


@app.get("/api/history")
def api_history():
    prod = request.args.get("product") or _S()["product"]
    return jsonify({"product": prod, "history": _history_list(prod)})


@app.get("/api/history/download")
def api_history_download():
    prod = request.args.get("product") or _S()["product"]
    entry = store.get_history(prod, request.args.get("id", ""))
    if not entry or not entry.report().exists():
        return jsonify({"ok": False, "error": "历史记录不存在。"}), 404
    return send_file(entry.report(), as_attachment=True,
                     download_name=f"{entry.title}_{entry.period}.docx")


def main():
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--host", default="127.0.0.1")
    ap.add_argument("--port", type=int, default=8000)
    args = ap.parse_args()
    print(f"KOL 复盘分析 · 本地 Web —— 引擎：{engine.label()}")
    print(f"产品：{'、'.join(f'{k}({v})' for k, v in PRODUCTS.items())}")
    print(f"打开浏览器： http://{args.host}:{args.port}")
    app.run(host=args.host, port=args.port, threaded=True)


if __name__ == "__main__":
    main()
