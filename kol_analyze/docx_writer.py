"""渲染复盘 .docx：
  一、广告部份（大盘/设计vsKOL/KOL分国家 + 概述）
  二、产出 vs 消耗 缺口分析（一句话总结 + 缺口表）—— 核心
  三、KOL 分语言素材分析（明细表）
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .metrics import Analysis
from .scripts import ScriptAnalysis

_HEADER_BG = "2F5496"
_ALT_BG = "F2F5FB"
_VERDICT_COLOR = {
    "加量": "2E7D32", "高潜": "1565C0", "维持": "555555",
    "削减": "C62828", "减少": "C62828", "覆盖缺口": "E65100",
    "待补全": "8E24AA",
}


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color}))


def _multiline(cell, text, bold_first=False, size=9, color=None):
    cell.text = ""
    para = cell.paragraphs[0]
    for i, line in enumerate(str(text).split("\n")):
        if i > 0:
            para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(line)
        run.font.size = Pt(size)
        if bold_first and i == 0:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _heading(doc, text, size, color="1F3864", bold=True, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def _body(doc, text, size=10, color=None):
    for line in str(text).split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)


def _mk_table(doc, headers, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = 1
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], _HEADER_BG)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
    return t, widths


def _apply_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]


def _pct(v):
    return f"{v:.2f}%" if v is not None else "—"


_STRAT_COLOR = {
    "维持精选": "555555", "探索新脚本": "1565C0",
    "挖新脚本": "E65100", "收窄精做": "C62828", "待补全": "8E24AA",
}


def _render_scripts(doc, data, sa: ScriptAnalysis):
    _heading(doc, "四、素材/脚本维度分析（跨语言）", 15, before=14)
    sec = data.get("script_section", {})
    if sec.get("overview"):
        _body(doc, sec["overview"])

    # 4.1 跨语言迁移建议
    _heading(doc, "4.1 跨语言脚本/形式迁移建议", 11.5, color="2F5496", before=8)
    migs = sec.get("migrations") or []
    for mg in migs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(mg).font.size = Pt(9.5)

    # 迁移数据表（脚本 × 跑出语言 × 建议扩展）
    if sa.migrations:
        tb, w = _mk_table(doc, ["脚本", "已跑出语言", "覆盖语言数", "建议扩展到"],
                          [Pt(110), Pt(90), Pt(70), Pt(150)])
        for r in sa.migrations[:12]:
            cells = tb.add_row().cells
            _multiline(cells[0], r.theme, bold_first=True, size=9)
            _multiline(cells[1], r.best_lang or "—", size=9)
            _multiline(cells[2], str(len(r.cells)), size=9)
            _multiline(cells[3], "、".join(r.migrate_to), size=9, color="1565C0")
        _apply_widths(tb, w)

    # 4.2 形式覆盖
    _heading(doc, "4.2 形式覆盖（口播 / 街采 / MV模板 …）", 11.5, color="2F5496", before=8)
    for fs in sec.get("format_suggestions") or []:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(fs).font.size = Pt(9.5)
    if sa.formats:
        tb, w = _mk_table(doc, ["形式", "在用语言（条数）", "建议试的语言"],
                          [Pt(70), Pt(180), Pt(140)])
        for f in sa.formats:
            cells = tb.add_row().cells
            _multiline(cells[0], f.fmt, bold_first=True, size=9)
            _multiline(cells[1], "、".join(f"{k}{v}" for k, v in
                                          sorted(f.present.items(), key=lambda x: -x[1])),
                       size=8.5)
            _multiline(cells[2], "、".join(f.suggest_to) or "—", size=9, color="1565C0")
        _apply_widths(tb, w)

    # 4.3 各语言脚本策略
    _heading(doc, "4.3 各语言脚本策略（继续 / 优化 / 砍 / 挖新脚本）", 11.5,
             color="2F5496", before=8)
    if sa.lang_strategies:
        tb, w = _mk_table(doc, ["语言", "脚本数", "跑出率", "档位", "建议"],
                          [Pt(60), Pt(46), Pt(50), Pt(64), Pt(200)])
        sug_by_name = {s["name"]: s["suggestion"]
                       for s in sec.get("lang_strategies", []) if s.get("name")}
        for s in sa.lang_strategies:
            cells = tb.add_row().cells
            _multiline(cells[0], s.name, bold_first=True, size=9)
            _multiline(cells[1], str(s.diversity), size=9)
            _multiline(cells[2], _pct(s.breakout), size=9)
            _multiline(cells[3], s.verdict, size=9,
                       color=_STRAT_COLOR.get(s.verdict, "555555"))
            _multiline(cells[4], sug_by_name.get(s.name, s.suggestion), size=9)
        _apply_widths(tb, w)


def render(data: dict, analysis: Analysis, scripts: ScriptAnalysis, out_path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10)

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(data.get("title", "月度 KOL 广告复盘"))
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("1F3864")
    per = doc.add_paragraph()
    per.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = per.add_run(data.get("period", ""))
    pr.font.size = Pt(13)
    pr.font.color.rgb = RGBColor.from_string("2F5496")

    m = analysis.market

    # ---- 一、广告部份 ----
    _heading(doc, "一、广告部份", 15)
    ad = data.get("ad_section", {})
    _body(doc, ad.get("overview", ""))
    if ad.get("caveat"):
        _heading(doc, "口径提醒", 10.5, color="C00000", before=6)
        _body(doc, ad["caveat"], size=9.5, color="C00000")

    if m.ad_country_share:
        _heading(doc, "广告大盘 · 分国家消耗（设计+KOL）", 11, color="2F5496", before=8)
        tb, w = _mk_table(doc, ["国家", "消耗", "占比"], [Pt(120), Pt(90), Pt(70)])
        for c, p in sorted(m.ad_country_share.items(), key=lambda x: -(x[1] or 0))[:15]:
            cells = tb.add_row().cells
            _multiline(cells[0], c, size=9)
            _multiline(cells[1], f"{m.ad_country_spend.get(c, '')}", size=9)
            _multiline(cells[2], f"{p:.2f}%", size=9)
        _apply_widths(tb, w)

    if m.kol_share_of_total is not None:
        _body(doc, f"设计师 vs KOL：KOL 占整体约 {m.kol_share_of_total:.2f}%。", size=10)

    # ---- 二、缺口分析 ----
    _heading(doc, "二、产出 vs 消耗 · 缺口分析", 15, before=14)
    if data.get("gap_summary"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("一句话总结：")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string("C00000")
        run2 = p.add_run(data["gap_summary"])
        run2.font.size = Pt(10.5)

    if analysis.gaps:
        headers = ["语言", "大盘消耗占比", "KOL消耗占比", "产出占比", "跑出率", "档位", "一句话结论"]
        widths = [Pt(56), Pt(70), Pt(64), Pt(56), Pt(50), Pt(48), Pt(150)]
        tb, w = _mk_table(doc, headers, widths)
        for g in analysis.gaps:
            cells = tb.add_row().cells
            _multiline(cells[0], g.name, bold_first=True, size=9)
            _multiline(cells[1], _pct(g.ad_market_share), size=9)
            _multiline(cells[2], _pct(g.kol_spend_share), size=9)
            _multiline(cells[3], _pct(g.publish_share), size=9)
            _multiline(cells[4], _pct(g.breakout_rate), size=9)
            _multiline(cells[5], g.verdict, size=9,
                       color=_VERDICT_COLOR.get(g.verdict, "555555"))
            _multiline(cells[6], g.one_line, size=8.5)
        _apply_widths(tb, w)

    # ---- 三、KOL 分语言素材分析 ----
    _heading(doc, "三、KOL 分语言素材分析", 15, before=14)
    headers = ["语言", "转化情况", "素材分析", "todo"]
    widths = [Pt(70), Pt(120), Pt(170), Pt(140)]
    tb, w = _mk_table(doc, headers, widths)
    for idx, lb in enumerate(data.get("langs", [])):
        cells = tb.add_row().cells
        name = lb.get("name", "")
        one = lb.get("one_liner", "")
        _multiline(cells[0], name + (f"\n{one}" if one else ""), bold_first=True)
        _multiline(cells[1], lb.get("conversion", ""))
        _multiline(cells[2], lb.get("creative_analysis", ""))
        _multiline(cells[3], lb.get("todo", ""))
        if idx % 2 == 1:
            for c in cells:
                _shade(c, _ALT_BG)
    _apply_widths(tb, w)

    # ---- 四、素材/脚本维度分析 ----
    _render_scripts(doc, data, scripts)

    # ---- 五、人力分工与调整建议 ----
    staff = data.get("staffing_section") or {}
    if staff.get("overview") or staff.get("people"):
        _heading(doc, "五、人力分工与调整建议", 15, before=14)
        if staff.get("overview"):
            _body(doc, staff["overview"])
        ppl = staff.get("people") or []
        if ppl:
            tb, w = _mk_table(doc, ["负责人", "调整建议"], [Pt(70), Pt(360)])
            for i, pr in enumerate(ppl):
                cells = tb.add_row().cells
                _multiline(cells[0], pr.get("person", ""), bold_first=True, size=9.5)
                _multiline(cells[1], pr.get("suggestion", ""), size=9.5)
                if i % 2 == 1:
                    for c in cells:
                        _shade(c, _ALT_BG)
            _apply_widths(tb, w)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(12)
    fr = foot.add_run("本报告由 KOL 复盘分析工具自动生成，请结合业务判断复核。")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("808080")

    out = Path(out_path)
    doc.save(out)
    return out
